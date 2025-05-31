import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pyautogui

# Function to create a timestamp with microseconds
def get_timestamp():
    # Get current time with microseconds
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # Strip last 3 digits
    return timestamp

# Function to take a screenshot
def take_screenshot():
    screenshot = pyautogui.screenshot()
    timestamp = get_timestamp()  # Get timestamp for the screenshot filename
    screenshot_filename = f'screenshots/screenshot_{timestamp}.png'  # Save path for the screenshot
    screenshot.save(screenshot_filename)  # Save the screenshot in the 'screenshots' folder
    return screenshot_filename  # Return the saved screenshot filename

# Function to create or update the Word document with a new screenshot
def update_word_document(doc, screenshot_filename):
    # Create 'screenshots' folder if it doesn't exist
    if not os.path.exists('screenshots'):
        os.makedirs('screenshots')

    # Check if the document exists
    if not os.path.exists('screenshots/screenshots_table.docx'):
        doc = Document()
        doc.add_heading('Unit Testing', 0)  # Added Unit Testing heading
        doc.add_paragraph(f'Date: 21 Maret 2025')  # Added Date
        # Add other header information if needed (Created by, Work Unit, etc.)
        doc.add_paragraph("")  # Empty paragraph for spacing

        # Create the table with 3 columns: (Column 1, Column 2, Screenshot in Column 3)
        table = doc.add_table(rows=1, cols=3)  # Three columns now
        table.style = 'Table Grid'
        # Define column headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column 1'
        hdr_cells[1].text = 'Column 2'
        hdr_cells[2].text = 'Screenshot'  # Image goes in Column 3
    else:
        doc = Document('screenshots/screenshots_table.docx')
        # Check if the document has the table, if not, create it
        if not doc.tables:
            doc.add_heading('Unit Testing', 0)
            doc.add_paragraph(f'Date: 21 Maret 2025')
            doc.add_paragraph("")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Column 1'
            hdr_cells[1].text = 'Column 2'
            hdr_cells[2].text = 'Screenshot'  # Image goes in Column 3
        else:
            table = doc.tables[0]  # Get the first table in the document

    # Add a new row with empty columns and the screenshot image in column 3
    row_cells = table.add_row().cells
    
    # Column 1 - Empty (or any data you want to add here)
    row_cells[0].text = ''  # Column 1 is left empty
    
    # Column 2 - Empty (or any data you want to add here)
    row_cells[1].text = ''  # Column 2 is left empty
    
    # Column 3 - Screenshot (image only)
    row_cells[2].paragraphs[0].clear()  # Clear any existing paragraph (to avoid unwanted text)
    row_cells[2].paragraphs[0].add_run().add_picture(screenshot_filename, width=Inches(2.0))  # Add the image in column 3

    # Save the Word document
    doc.save('screenshots/screenshots_table.docx')
    print(f"Word document updated and saved as 'screenshots/screenshots_table.docx'.")

# Example: Call this function each time you want to take a screenshot and update the document
def capture_and_update():
    screenshot_filename = take_screenshot()
    update_word_document(None, screenshot_filename)  # None passed because we will handle document creation/updating internally

# Call the function as needed
# Each time you want to take a screenshot and update the document, call `capture_and_update()`
# For example, calling it multiple times will append rows with new screenshots
capture_and_update()  # Call this function whenever you need a screenshot and document update.
